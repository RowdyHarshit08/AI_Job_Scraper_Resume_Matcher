import os
import re
import gc
import fitz
from docx import Document


SKILLS = [
    "python",
    "java",
    "c++",
    "c",
    "sql",
    "oracle",
    "html",
    "css",
    "javascript",
    "django",
    "mongodb",
    "machine learning",
    "artificial intelligence",
    "ai",
    "data structures",
    "algorithms",
    "nlp",
    "git",
    "flask",
    "pandas",
    "numpy",
    "scikit-learn",
    "tensorflow",
    "pytorch",
    "excel",
    "statistics",
    "etl",
    "cloud",
    "data engineering",
    "rest api",
]


def extract_pdf_text(path):
    """
    First tries normal PDF text extraction.
    If the PDF is scanned/image-based, it automatically
    falls back to OCR page by page.
    """

    text = ""

    try:
        document = fitz.open(path)
    except Exception as e:
        return "", f"PDF could not be opened: {str(e)}"

    # Password-protected PDF
    if document.needs_pass:
        document.close()
        return "", "This PDF is password protected. Please upload an unlocked PDF."

    try:
        # -------------------------------------------------
        # STEP 1: Normal text extraction
        # -------------------------------------------------
        for page in document:
            try:
                page_text = page.get_text("text", sort=True)
                if page_text:
                    text += page_text + "\n"
            except Exception:
                continue

        text = text.strip()

        # If enough text was found, OCR is not required.
        if len(text) >= 50:
            document.close()
            return text, ""

        # -------------------------------------------------
        # STEP 2: OCR for scanned/image PDFs
        # -------------------------------------------------
        try:
            import pytesseract
            from PIL import Image
        except Exception as e:
            document.close()
            return "", f"OCR libraries are unavailable: {str(e)}"

        ocr_text = ""

        # Try Hindi + English first.
        # If Hindi language data is unavailable, English is used.
        try:
            languages = pytesseract.get_languages(config="")
        except Exception:
            languages = ["eng"]

        if "hin" in languages and "eng" in languages:
            ocr_lang = "eng+hin"
        elif "eng" in languages:
            ocr_lang = "eng"
        elif "hin" in languages:
            ocr_lang = "hin"
        else:
            ocr_lang = None

        if not ocr_lang:
            document.close()
            return "", "OCR language data is not installed on the server."

        for page_number, page in enumerate(document):

            try:
                # Controlled resolution to avoid Render memory problems.
                matrix = fitz.Matrix(1.5, 1.5)

                pix = page.get_pixmap(
                    matrix=matrix,
                    colorspace=fitz.csRGB,
                    alpha=False
                )

                image = Image.frombytes(
                    "RGB",
                    [pix.width, pix.height],
                    pix.samples
                )

                try:
                    page_ocr = pytesseract.image_to_string(
                        image,
                        lang=ocr_lang,
                        config="--psm 6",
                        timeout=30
                    )

                    if page_ocr:
                        ocr_text += page_ocr + "\n"

                except Exception:
                    # Try English as a final fallback
                    try:
                        page_ocr = pytesseract.image_to_string(
                            image,
                            lang="eng",
                            config="--psm 6",
                            timeout=30
                        )

                        if page_ocr:
                            ocr_text += page_ocr + "\n"

                    except Exception:
                        pass

                # Release memory after EVERY page.
                image.close()
                del pix
                gc.collect()

            except Exception:
                continue

        document.close()

        ocr_text = ocr_text.strip()

        if len(ocr_text) >= 20:
            return ocr_text, ""

        return "", (
            "Could not read text from this PDF. "
            "The PDF may be corrupted, encrypted, or the scan quality "
            "may be too poor for OCR."
        )

    except Exception as e:
        try:
            document.close()
        except Exception:
            pass

        return "", f"PDF processing failed: {str(e)}"


def extract_docx_text(path):
    try:
        document = Document(path)

        parts = []

        # Paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())

        # Tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        return "\n".join(parts)

    except Exception:
        return ""


def extract_txt_text(path):
    try:
        with open(
            path,
            "r",
            encoding="utf-8",
            errors="ignore"
        ) as file:
            return file.read()

    except Exception:
        return ""


def detect_skills(text):
    """
    Detect skills from extracted resume text.
    """

    lower_text = text.lower()

    detected = set()

    for skill in SKILLS:

        # Flexible matching for common variations.
        if skill.lower() in lower_text:
            detected.add(skill)

    return sorted(detected)


def extract_resume(path):
    """
    Main resume extraction function.

    Supports:
    PDF
    DOCX
    TXT
    """

    if not path or not os.path.exists(path):
        return "", {
            "error": "Resume file could not be found."
        }

    extension = os.path.splitext(path)[1].lower()

    text = ""
    error = ""

    # -------------------------------------------------
    # PDF
    # -------------------------------------------------
    if extension == ".pdf":

        text, error = extract_pdf_text(path)

    # -------------------------------------------------
    # DOCX
    # -------------------------------------------------
    elif extension == ".docx":

        text = extract_docx_text(path)

        if not text.strip():
            error = "Could not extract readable text from this DOCX file."

    # -------------------------------------------------
    # TXT
    # -------------------------------------------------
    elif extension == ".txt":

        text = extract_txt_text(path)

        if not text.strip():
            error = "Could not read this TXT file."

    # -------------------------------------------------
    # Unsupported file
    # -------------------------------------------------
    else:

        return "", {
            "error": (
                "Unsupported file type. "
                "Please upload a PDF, DOCX or TXT file."
            )
        }

    text = text.strip()

    if not text:

        return "", {
            "error": error or (
                "No readable text was found in this resume."
            )
        }

    # -------------------------------------------------
    # Extract candidate information
    # -------------------------------------------------

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    # Candidate name
    candidate_name = "Candidate"

    if lines:
        candidate_name = lines[0][:100]

    # Email
    email_match = re.search(
        r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}",
        text
    )

    email = (
        email_match.group(0)
        if email_match
        else ""
    )

    # Indian phone number
    phone_match = re.search(
        r"(?:\+91[\s-]?)?[6-9]\d{9}",
        text
    )

    phone = (
        phone_match.group(0)
        if phone_match
        else ""
    )

    # Skills
    skills = detect_skills(text)

    return text, {
        "name": candidate_name,
        "email": email,
        "phone": phone,
        "skills": skills,
    }
